from __future__ import annotations

from pathlib import Path
from uuid import uuid4
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.schemas.notes import TopicNote


def _safe_join_lines(lines: list[str]) -> str:
    return "\n".join(f"- {item}" for item in lines)


def write_markdown(export_dir: Path, notes: list[TopicNote]) -> Path:
    path = export_dir / f"{uuid4().hex}.md"
    sections: list[str] = []

    for topic in notes:
        sections.append(f"## {topic.topic_name}")
        sections.append("\n### 1. Explanation (What YouTuber Said)")
        sections.extend([f"- {line}" for line in topic.explanation])

        sections.append("\n### 2. Screen Content")
        sections.extend([f"- {line}" for line in topic.screen_content])
        if topic.formulas_or_diagrams:
            sections.append("\n**Formulas / Diagrams**")
            sections.extend([f"- {line}" for line in topic.formulas_or_diagrams])
        if topic.diagram:
            sections.append("\n**Topic Diagram**")
            sections.append(f"- {topic.diagram}")

        sections.append("\n### 3. Code Section")
        if not topic.code_sections:
            sections.append("- No code in this topic.")
        for block in topic.code_sections:
            sections.append(f"\nLanguage: `{block.language}`")
            sections.append("```" + block.language)
            sections.append(block.code)
            sections.append("```")
            sections.append(f"- {block.explanation}")
            for line in block.line_by_line:
                sections.append(f"  - Line {line.line_number}: {line.explanation}")

        sections.append("\n---\n")

    path.write_text("\n".join(sections), encoding="utf-8")
    return path


def write_docx(export_dir: Path, notes: list[TopicNote]) -> Path:
    path = export_dir / f"{uuid4().hex}.docx"
    doc = Document()
    doc.add_heading("Engineering Notes", 0)

    for topic in notes:
        doc.add_heading(topic.topic_name, level=1)
        doc.add_heading("1. Explanation (What YouTuber Said)", level=2)
        for item in topic.explanation:
            doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("2. Screen Content", level=2)
        for item in topic.screen_content:
            doc.add_paragraph(item, style="List Bullet")
        for item in topic.formulas_or_diagrams:
            doc.add_paragraph(f"Formula/Diagram: {item}", style="List Bullet")
        if topic.diagram:
            doc.add_paragraph(f"Topic Diagram: {topic.diagram}", style="List Bullet")

        doc.add_heading("3. Code Section", level=2)
        if not topic.code_sections:
            doc.add_paragraph("No code in this topic.", style="List Bullet")

        for block in topic.code_sections:
            doc.add_paragraph(f"Language: {block.language}")
            doc.add_paragraph(block.code)
            doc.add_paragraph(f"Explanation: {block.explanation}", style="List Bullet")
            for line in block.line_by_line:
                doc.add_paragraph(f"Line {line.line_number}: {line.explanation}", style="List Bullet")

    doc.save(path)
    return path


def write_pdf(export_dir: Path, notes: list[TopicNote]) -> Path:
    path = export_dir / f"{uuid4().hex}.pdf"
    styles = getSampleStyleSheet()
    story = [Paragraph("Engineering Notes", styles["Title"]), Spacer(1, 12)]

    for topic in notes:
        story.append(Paragraph(topic.topic_name, styles["Heading2"]))
        story.append(Paragraph("1. Explanation (What YouTuber Said)", styles["Heading3"]))
        story.append(Paragraph(_safe_join_lines(topic.explanation).replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 8))

        story.append(Paragraph("2. Screen Content", styles["Heading3"]))
        story.append(Paragraph(_safe_join_lines(topic.screen_content).replace("\n", "<br/>"), styles["BodyText"]))
        if topic.formulas_or_diagrams:
            story.append(Paragraph(_safe_join_lines(topic.formulas_or_diagrams).replace("\n", "<br/>"), styles["BodyText"]))
        if topic.diagram:
            story.append(Paragraph(f"- {topic.diagram}", styles["BodyText"]))

        story.append(Spacer(1, 8))
        story.append(Paragraph("3. Code Section", styles["Heading3"]))
        if not topic.code_sections:
            story.append(Paragraph("- No code in this topic.", styles["BodyText"]))

        for block in topic.code_sections:
            story.append(Paragraph(f"Language: {block.language}", styles["BodyText"]))
            story.append(Paragraph(block.code.replace("\n", "<br/>"), styles["BodyText"]))
            story.append(Paragraph(block.explanation, styles["BodyText"]))
            for line in block.line_by_line:
                story.append(Paragraph(f"Line {line.line_number}: {line.explanation}", styles["BodyText"]))

        story.append(Spacer(1, 12))

    doc = SimpleDocTemplate(str(path), pagesize=A4)
    doc.build(story)
    return path
